from pathlib import Path
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
for path in sorted(Path("output").glob("*.docx")):
    doc = Document(path)
    root = ET.fromstring(zipfile.ZipFile(path).read("word/document.xml"))
    widths = [int(x.attrib.get("{%s}w" % NS["w"], "0")) for x in root.findall(".//w:tblW", NS)]
    text = "\n".join(p.text or "" for p in doc.paragraphs)
    print(path.name)
    print(" paragraphs", len(doc.paragraphs), "tables", len(doc.tables))
    print(" headings", sum(p.style.name == "Heading 1" for p in doc.paragraphs), sum(p.style.name == "Heading 2" for p in doc.paragraphs), sum(p.style.name == "Heading 3" for p in doc.paragraphs))
    print(" page", doc.sections[0].page_width, doc.sections[0].page_height, "margins", doc.sections[0].left_margin, doc.sections[0].right_margin)
    print(" table widths", sorted(set(widths)))
    print(" bad tokens", [x for x in ("TODO", "TBD", "placeholder") if x in text])
