from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
OUT.mkdir(exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RISK = "9B1C1C"
GOLD = "7A5A00"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, color=None, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc, title_short):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 25, INK, 0, 6),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=name.startswith("Heading"))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    set_style_font(code, name="Courier New", size=8.5, color=INK)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0

    if "Small" not in styles:
        small = styles.add_style("Small", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = styles["Small"]
    set_style_font(small, size=9, color=MUTED)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.1

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = "县域生活信息交换服务平台 | " + title_short
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "内部评审稿 | 第 "
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)
    add_page_number(footer)
    return doc


def add_title_block(doc, title, subtitle, code, version="V1.1 A/B/C 融合评审稿"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, size=25, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, size=13, color=MUTED)
    rows = [
        ("文档编号", code),
        ("版本", version),
        ("日期", "2026-09-02"),
        ("依据", "spec.docx、config_matrix.md"),
        ("适用范围", "单县单实例 MVP，供产品、设计、开发、运营评审"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for c in cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=9.5, color=INK)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    set_table_geometry(table, [1800, 7560])
    doc.add_paragraph()


def add_para(doc, text, bold_prefix=None, style=None, color=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=10.5, color=color or INK, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=10.5, color=color or INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=color or INK)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.2, color=INK)


def add_callout(doc, label, text, fill=CALLOUT, color=INK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    for i, line in enumerate(text.strip("\n").split("\n")):
        r = p.add_run(line)
        set_run_font(r, name="Courier New", size=8.5, color=INK)
        if i < len(text.strip("\n").split("\n")) - 1:
            r.add_break()
    return p


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=font_size, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=font_size, color=INK)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def source_doc():
    from docx import Document as SourceDocument
    return SourceDocument(ROOT / "spec.docx")


def table_rows(src, index):
    table = src.tables[index]
    return [[c.text.replace("\n", " / ").strip() for c in row.cells] for row in table.rows]


FR_TABLES = [
    (7, "账号与身份"), (8, "位置与圈层"), (9, "语音输入"), (10, "语音输出"),
    (11, "一键发布"), (12, "有效期"), (14, "审核"), (17, "逛一逛"),
    (18, "搜索"), (19, "AI 对话"), (21, "助农供求"), (22, "求职招工与有偿任务"),
    (23, "二手市场"), (24, "约局互动"), (25, "消息与私聊"), (26, "我的"),
    (28, "公众号"), (31, "后台管理端"), (33, "反诈发布侧"), (34, "反诈浏览侧"),
    (35, "反诈私聊侧"), (36, "反诈处置侧"), (40, "配置中心"), (52, "AI 双出口补充"),
]


DOMAIN_LOGIC = {
    "ACC": ("登录、注册或触碰联系方式", "登录弹层、身份补全页、脱敏联系方式", "用户/权限状态更新，联系方式访问写审计日志"),
    "LOC": ("首次进入、授权位置或切换范围", "乡镇选择器、范围胶囊、相对距离标签", "保存县/镇/村行政区划码，拒绝授权时不阻塞"),
    "VOI": ("按住录音、松手上传或识别失败", "录音浮层、识别状态、可编辑文本", "语音 OSS 直传，ASR/补全进入 Redis Stream，失败降级纯语音"),
    "PUB": ("点击发布、确认草稿或续期", "发布三屏、草稿恢复、有效期选择", "Post 状态由 draft -> pending/published，审核和到期任务异步处理"),
    "AUDIT": ("发布、机审命中或举报", "审核状态、原因、人工队列", "风险等级驱动 pending/rejected/published，所有动作留痕"),
    "BRW": ("进入逛一逛、下拉刷新或翻页", "双列卡片、标签、骨架屏", "按状态/类目/乡镇过滤，游标分页，首屏缓存 60 秒"),
    "SCH": ("输入关键词或按住语音搜索", "搜索框、筛选、无结果兜底", "MySQL ngram + 同义词，零结果转求购草稿"),
    "AI": ("输入文字/语音或点预置问题", "流式消息、结果卡片、快捷动作", "意图路由 -> 结构化过滤 -> 召回/发布草稿，模型不可用降级搜索"),
    "AGR": ("发布或筛选农产品供求", "我要卖/我要收标签、品种和供需字段", "品种词库归一，上市窗口驱动 valid_until"),
    "JOB": ("发布招工、报名或任务到期", "三要素校验、还缺人数、报名按钮", "报名原子扣减名额，联系方式互见，完成状态进入 deal"),
    "SEC": ("发布二手或标记已卖出", "首图、成色、议价、同城自提提示", "已售状态置灰保留，风险提示常驻"),
    "ACT": ("发起约局、报名、取消或满员", "类型、时间、地点、费用方式", "满员自动关闭，打牌屏蔽金额并强制承诺不赌博"),
    "MSG": ("详情点击在线聊或发送消息", "会话页、未读红点、反诈卡", "长轮询拉取消息，命中词/二维码插入不可关闭提示"),
    "MINE": ("进入我的、查看发布/报名/收藏", "按状态分组、续期、帮助中心", "仅展示本人可见状态，续期刷新 bumped_at"),
    "MP": ("关注公众号、点击菜单或接收推送", "菜单、图文模板、召回卡片", "公众号只做触达，复杂动作跳小程序，周频控 3 条"),
    "ADM": ("审核、举报、用户、类目、配置或指标管理", "PC 管理台列表、详情抽屉、操作确认", "同一 post 表和状态机，角色枚举中间件，导出留痕"),
    "ANTI": ("发布、浏览详情、私聊、举报或确认诈骗", "类目警示条、反诈卡、举报表单", "风险词和行为风控实时拦截，封禁维度 openid+手机号+设备号"),
}


def fr_detail(fid, need):
    prefix = fid.split("-")[1] if "-" in fid else ""
    trigger, ui, backend = DOMAIN_LOGIC.get(prefix, ("用户执行对应动作", "对应页面控件和状态反馈", "服务端校验、记录日志并更新状态"))
    acceptance = "可重复验证：触发条件满足时行为正确；失败时有明确提示；刷新页面后状态与服务端一致。"
    if fid.endswith("01"):
        acceptance = "首条主路径可在真实微信环境完成；游客被拦截点明确；异常网络不丢草稿。"
    if prefix == "PUB":
        acceptance = "发布成功后返回 post_id；低风险按目标时延上线；待审/驳回原因可见；重复提交不产生重复 Post。"
    if prefix == "AUDIT":
        acceptance = "审核员能在一屏看到判断依据；处置结果写入审计日志并触发对应通知。"
    if prefix == "AI":
        acceptance = "结果卡片字段全部来自数据库；最多一个澄清问题；模型不可用时仍可搜索或进入发布。"
    if prefix == "ANTI":
        acceptance = "命中规则立即出现对应拦截/警示；举报、封禁和通知链路可追溯；误封存在申诉入口。"
    return f"触发：{trigger}。页面：{ui}。实现：{backend}。验收：{acceptance}"


def add_growth_requirements(doc):
    doc.add_heading("A/B/C 信息交换增长闭环", level=2)
    add_para(doc, "A/B/C 不是单独的营销模块，而是把查找、发布和回访串成一条可度量的信息交换链：用户先能快速找到，再能低成本发布，最后因为真实回应和本镇动态愿意回来。首期只使用真实搜索、发布、响应和订阅事件，不引入虚假浏览量、复杂积分或红包裂变。")
    add_table(doc, ["方向", "用户动机", "首期功能要求", "核心指标"], [
        ("A｜让人想查", "我输入一句话，就能找到本镇有用信息", "首页可见搜索框；热门/急需搜索词；自然语言搜索；结果卡片显示距离、时效、响应状态；无结果转求助/求购草稿", "搜索发起率 >= 35%；搜索后详情点击率 >= 45%；零结果转发布率 >= 15%"),
        ("B｜让人想发", "我发出去，很快有人看到或回应", "语音/文字自动结构化；发布前展示覆盖范围；发布后显示浏览、收藏、响应；有人回应时通知；支持补图、续期和状态更新", "发布完成率 >= 60%；24 小时首次响应率 >= 50%；发布后响应率 >= 30%"),
        ("C｜让人常回来", "每天打开都能看到镇上新发生的事", "本镇脉搏；刚刚/急需/即将截止标签；每日信息摘要；订阅提醒；已找到、已招满、已完成等真实状态", "次周留存 >= 25%；摘要打开率 >= 20%；过期信息续期率 >= 10%"),
    ], [1250, 2050, 3820, 2240], font_size=8.1)
    add_bullets(doc, [
        "所有激励必须来自真实行为：搜索、详情点击、收藏、联系、响应、发布完成、状态更新和订阅打开。",
        "用户没有查到时不能停在‘暂无结果’，必须给出相近结果、订阅提醒或一键生成发布草稿。",
        "用户发布后不能只显示浏览量，要明确告诉他是否有人收藏、联系、报名或已完成。",
        "首页优先展示本镇新鲜、急需、即将截止的信息，避免做成泛内容社区。",
    ])


def add_growth_feature_matrix(doc):
    doc.add_heading("A/B/C 增长闭环功能点", level=2)
    add_para(doc, "以下条目补充源需求中的搜索、发布、消息和运营能力，编号可直接拆成开发任务。")
    add_table(doc, ["编号", "功能要求", "优先级", "交互 / 状态 / 验收口径"], [
        ("FR-GRO-01", "首页搜索与热门示例词", "P0", "首页首屏展示输入框和 3-5 个本镇热搜词；可输入‘收玉米/找零工/借农机’等自然语言；回车或点击搜索进入结果；记录 search_submitted 事件。"),
        ("FR-GRO-02", "结果解释与真实热度标签", "P0", "结果卡片显示同镇/邻镇、相对时间、距离、响应状态和‘刚刚/急需/即将截止’标签；标签必须由真实事件和配置阈值生成，不能手工造数。"),
        ("FR-GRO-03", "无结果转发布草稿", "P0", "无结果时展示相近词、订阅提醒和‘帮我发布’；AI/规则将 query 预填为标题或正文草稿；用户确认前不得自动发布。"),
        ("FR-GRO-04", "发布后的响应闭环", "P0", "发布成功页和‘我的’展示浏览、收藏、联系、报名、响应人数；首次响应触发站内消息/订阅通知；发布者可标记已找到/已招满/已完成。"),
        ("FR-GRO-05", "本镇脉搏与信息摘要", "P0", "首页展示近 24 小时新信息、招工、响应和新发布；后台按乡镇生成每日摘要；数据不足时显示真实的‘今日暂无’和发布入口。"),
        ("FR-GRO-06", "订阅与到期召回", "P1", "用户可订阅关键词+乡镇+类目；新匹配和到期前提醒可退订；每天最多一次同类提醒，写入推送结果和打开事件。"),
        ("FR-GRO-07", "增长事件与指标看板", "P0", "统一记录 search_submitted、search_zero_result、draft_created、post_published、response_received、post_completed、digest_opened；后台按乡镇/类目/日期查询漏斗。"),
    ], [1150, 2250, 700, 5260], font_size=8.2)


def add_growth_architecture(doc):
    doc.add_heading("A/B/C 增长闭环设计", level=2)
    add_para(doc, "采用‘查找 -> 发布 -> 回应 -> 回访’四段式闭环。A 负责降低查找成本，B 负责让发布后有真实回音，C 负责把本镇新鲜信息和到期提醒变成回访理由。所有事件先写 events 表，再由通知、摘要和指标任务异步消费。")
    add_table(doc, ["闭环阶段", "客户端", "服务端 / AI", "单人实现边界"], [
        ("A 查找", "首页搜索、热搜词、语音搜索、结果解释、相近词", "规则/同义词 -> MySQL ngram -> 可选向量召回；结果附距离、时效、响应状态", "先做 MySQL ngram + 词库，不上独立搜索集群"),
        ("B 发布", "长按说话/输入 -> 草稿预览 -> 覆盖范围 -> 发布", "ASR、字段抽取、风险审核、草稿缺失字段和 warnings", "最多三次主动交互；模型不可用时允许纯语音/手填"),
        ("B 回应", "浏览/收藏/联系/报名反馈，已找到/已招满/已完成", "response 事件、去重、通知和状态机，禁止客户端自报数据", "先做站内消息，订阅通知做 P1"),
        ("C 回访", "本镇脉搏、今日摘要、订阅、到期续期", "按乡镇聚合真实事件，定时生成摘要并限频推送", "单乡镇定时任务即可，不做复杂推荐系统"),
    ], [1300, 2850, 3150, 2060], font_size=8.2)
    add_table(doc, ["指标", "MVP 目标", "事件来源", "看板维度"], [
        ("搜索发起率", ">= 35%", "search_submitted / home_view", "乡镇、入口、日期"),
        ("零结果转发布率", ">= 15%", "draft_created / search_zero_result", "query、类目、乡镇"),
        ("发布完成率", ">= 60%", "post_published / draft_created", "发布入口、类目、用户类型"),
        ("24 小时首次响应率", ">= 50%", "response_received / post_published", "类目、乡镇、信息完整度"),
        ("次周留存", ">= 25%", "week2_active / week1_new_user", "来源、乡镇、首个动作"),
    ], [1900, 1300, 3200, 2960], font_size=8.2)
    add_code(doc, """
search_submitted -> search_zero_result? -> search_result_clicked
       |                                      |
       └-> subscribe_created / draft_created  └-> contact / favorite / signup
                                                   |
post_published -> response_received -> post_completed
       |
       └-> digest_generated / digest_opened / expiry_reminded / post_renewed
""")
    add_bullets(doc, [
        "搜索零结果必须落一条可操作路径：相近词、订阅提醒或预填发布草稿，不能只返回空数组。",
        "发布成功页首屏展示真实浏览/收藏/联系/报名计数，并提供补图、补地点、续期和状态更新。",
        "响应通知按 post_id 去重和限频；同一用户同一条信息每天最多收到一条提醒。",
        "每日摘要只聚合已发布、未过期信息；数据不足时明确显示‘今日暂无’，不伪造热度。",
    ])


def add_requirements_matrix(doc, src):
    doc.add_heading("五、功能点明细", level=1)
    add_para(doc, "以下矩阵把源需求中的 FR 编号转为可执行条目。每条都同时描述触发、页面、状态与验收，开发时可直接拆成任务卡。P0 为 MVP 必须，P1 为 MVP 尽量，P2 为二期。")
    for idx, domain in FR_TABLES:
        rows = table_rows(src, idx)
        if not rows or len(rows[0]) < 2:
            continue
        doc.add_heading(domain, level=2)
        header = rows[0]
        data = rows[1:]
        out = []
        for row in data:
            if len(row) >= 3:
                fid, need, priority = row[0], row[1], row[2]
            else:
                fid, need, priority = row[0], row[1], "P0"
            out.append((fid, need, priority, fr_detail(fid, need)))
        add_table(doc, ["编号", "功能要求", "优先级", "交互 / 状态 / 验收口径"], out, [1150, 2900, 800, 4510], font_size=8.2)


def add_function_doc():
    src = source_doc()
    doc = configure_document(Document(), "功能点明细")
    add_title_block(doc, "县域生活信息交换服务平台", "功能点明细与低保真页面解析", "F-REQ-20260902")
    add_callout(doc, "结论先行", "首期不是做一个内容社区，而是做一条可闭环的县域信息撮合链：用户能说、系统能结构化、信息能被看见、有人能联系、风险能被处置。")
    doc.add_heading("阅读导航", level=1)
    add_bullets(doc, [
        "第一至四章：产品边界、角色、信息架构和共用状态机。",
        "第五章：按 FR 编号展开的功能点明细，可直接转开发任务。",
        "第六章：低保真页面和关键流程的页面级解析。",
        "第七至八章：全局状态、验收清单和待确认事项。",
    ])
    doc.add_heading("一、产品目标与边界", level=1)
    add_para(doc, "产品定位：一个县域范围内、以语音和一键发布为主要输入方式的本地生活信息撮合平台。四个业务模块共用一套 Post 模型，平台只提供信息发布、展示、搜索、沟通和风险治理，不介入资金托管、担保交易、物流履约或交易结算。")
    add_table(doc, ["MVP 北极星指标", "目标", "对应产品动作"], [
        ("首次响应率", "发布后 24 小时内有人联系 >= 50%", "默认乡镇范围、时效排序、结果卡片一键联系"),
        ("发布转化率", "进入发布页的用户 >= 60% 完成发布", "语音优先、最多三次交互、未识别字段留空不阻断"),
        ("语音发布占比", ">= 40%", "长按录音、ASR 异步、纯语音降级"),
        ("搜索发起率", "首页访问用户中 >= 35% 发起搜索", "可见搜索框、热门词、语音搜索和搜索结果解释"),
        ("零结果转发布率", "无结果搜索中 >= 15% 进入发布草稿", "相近词、订阅提醒、query 预填发布草稿"),
        ("发布后响应率", "发布后 24 小时内收到收藏/联系/报名 >= 30%", "响应事件、状态更新、站内消息和订阅通知"),
        ("次周留存", ">= 25%", "到期续期、公众号召回、订阅提醒"),
        ("单条信息成本", "<= 0.02 元", "规则优先、模型分层、TTS 缓存"),
    ], [1800, 2100, 5460])
    add_growth_requirements(doc)
    add_callout(doc, "明确不做", "首期不做资金托管、物流、信用分、关注点赞、个性化信息流、跨县多租户和复杂 RBAC。所有后续能力只做数据结构、接口和 Feature Flag 预留。", fill="FFF8E8", color=GOLD)

    doc.add_heading("二、用户与角色", level=1)
    add_table(doc, ["角色", "典型诉求", "设计约束", "首期关键权限"], [
        ("种植户/养殖户", "卖货、找收购、打字困难", "大字号、语音优先、三步内", "浏览、语音发布、查看联系方式、私聊、举报"),
        ("返乡/在县青壮年", "找活、买卖闲置、约局", "需要可逛的信息流", "浏览、搜索、发布、报名、收藏、私聊"),
        ("小微雇主/商户", "临时快速找人", "高频、急、强调人数和薪资", "招工、任务发布、报名处理、续期"),
        ("村级信息员", "代他人发布", "多条信息、代发留痕", "代发、查看代发质量、举报"),
        ("运营/审核", "治理、审核、撮合", "30 分钟处理一天队列", "审核、下架、封禁、配置、看板、导出"),
    ], [1550, 2200, 2300, 3310])
    doc.add_heading("权限矩阵", level=2)
    add_table(doc, ["能力", "游客", "注册用户", "认证发布者", "村级信息员", "运营/审核"], [
        ("浏览列表与详情", "是", "是", "是", "是", "是"),
        ("查看联系方式", "否", "是", "是", "是", "是"),
        ("发布信息", "否", "限额", "限额更高", "是", "是"),
        ("代他人发布", "否", "否", "否", "是", "是"),
        ("审核/下架/封禁", "否", "否", "否", "否", "是"),
    ], [2100, 1050, 1500, 1500, 1500, 1710])

    doc.add_heading("三、信息架构与共用规则", level=1)
    add_code(doc, """
小程序 / H5
├─ 首页：AI 语音入口 + 免责反诈条 + 九宫格 + 首屏信息
├─ 逛一逛：双列瀑布流 + 类目胶囊 + 搜索/筛选
├─ 发布：长按录音 -> AI 草稿 -> 确认 -> 审核结果
├─ 消息：会话、系统通知、报名/审核/到期提醒
└─ 我的：我的发布、报名、收藏、帮助中心、防骗指南

全局入口：右下角 AI 悬浮球；详情页主操作为在线聊，电话为次操作。
公众号：资讯、召回、周报、AI 轻问答；复杂动作统一跳小程序。
后台：审核、举报、用户、类目、配置、运营看板、导出。
""")
    doc.add_heading("统一 Post 模型", level=2)
    add_table(doc, ["字段组", "字段", "说明"], [
        ("基础", "id, author_id, category, sub_category", "四模块共用，类目差异放 ext JSON"),
        ("内容", "title, body_text, raw_asr_text, raw_voice_url", "保留原始输入，支持重跑模型与审计"),
        ("交易信息", "price, price_unit, quantity, quantity_unit", "可空；约局打牌类默认禁用金额"),
        ("位置", "location_code, lng, lat", "主检索使用县/镇/村码，坐标仅用于地图跳转"),
        ("联系", "contact_type, contact_value", "默认脱敏，完整查看写访问日志"),
        ("状态", "draft/pending/published/closed/expired/removed", "所有端共用状态机"),
        ("时效", "valid_until, bumped_at", "到期自动下架；续期刷新 bumped_at"),
        ("排序", "quality_score, weight_boost", "首期 weight_boost=1，置顶能力预留"),
    ], [1450, 3600, 4310])
    doc.add_heading("共用状态机", level=2)
    add_code(doc, """
draft -> pending_review -> published -> closed / expired / removed
                 └──────-> rejected -> draft（用户修改后重提）
published --举报/风控--> pending_review
expired --一键续期--> published（valid_until 延长，bumped_at 刷新）
""")
    add_para(doc, "状态显示必须区分本人视图和他人视图：作者可见待审、驳回原因和申诉入口；其他用户只看到已发布信息。")

    add_requirements_matrix(doc, src)
    add_growth_feature_matrix(doc)

    doc.add_heading("六、低保真页面与流程解析", level=1)
    add_para(doc, "线框采用文字低保真，重点不是视觉稿，而是确认首屏信息、动作顺序、异常状态和页面之间的跳转关系。实现时先按这些骨架做可点击原型，再补视觉 token。")
    wireframes = [
        ("6.1 首页", """
┌────────────────────────────────┐
│ 县域生活     城关镇 ▾     消息(2) │
│ ┌────────────────────────────┐ │
│ │ 说句话，帮你找或发信息       │ │
│ │        [按住说话]             │ │
│ │       [改用文字输入]          │ │
│ └────────────────────────────┘ │
│ 平台不参与交易，不担保资金       │
│ 本镇脉搏：28 条新信息 · 4 人响应 │
│ [助农] [招工] [任务]             │
│ [二手] [约局] [全部]             │
│ 本镇最新                         │
│ [收] 城关镇收玉米  1.2 元/斤     │
│ [招] 下午装车缺 5 人  260 元/天  │
├────首页──逛一逛──  发布  ─消息─我的┤
└────────────────────────────────┘
""", ["首屏必须同时出现 AI 入口和免责/反诈条。", "类目用图标 + 大字，发布按钮居中凸起。", "卡片显示收/招/局标签、同镇或相对距离。"]),
        ("6.2 逛一逛", """
┌────────────────────────────────┐
│ 逛一逛 [收玉米、找零工……] [搜索] │
│ [全部][助农][招工][任务][二手][约局]│
│ 本镇热搜：[收鸡蛋] [下午装车]     │
│ 综合 ▾   同镇 ▾                  │
│ ┌──────────┐ ┌──────────┐      │
│ │ [图/色块] │ │ [图/色块] │      │
│ │ 旧打谷机  │ │ 招木工2人 │      │
│ │ 1200元    │ │ 260元/天  │      │
│ │ 同镇 · 2h │ │ 邻镇 · 1h │      │
│ └──────────┘ └──────────┘      │
│ 无结果：换个说法 / 订阅提醒 / 帮我发布 │
│            [加载更多]            │
└────────────────────────────────┘
""", ["无图信息使用类目色块，不留空白洞。", "默认综合=距离×时新度×完整度，可切最新/最近。", "使用游标分页，列表接口不返回正文和联系方式。"]),
        ("6.3 一键发布三屏", """
屏1 说话：
┌────────────────────────────────┐
│ 发布信息                         │
│          [长按说话]              │
│ 也可以 [改用文字输入] 或上传图片   │
│ 平台不参与交易，不担保资金         │
└────────────────────────────────┘

屏1B 文字输入：
┌────────────────────────────────┐
│ 写下你想找或想发布的内容           │
│ ┌────────────────────────────┐ │
│ │ 城关镇有没有人收玉米？       │ │
│ │                            │ │
│ └────────────────────────────┘ │
│             [开始查找]          │
│ [切回语音]                      │
└────────────────────────────────┘

屏2 核对一下：
┌────────────────────────────────┐
│ ‹ 返回       核对一下       发布 │
│ [原始语音 00:12]   [改文字]      │
│ 分类 助农 > 农产品出售  [改]     │
│ 标题 城关镇土鸡蛋 200斤          │
│ 价格 1.20 元/斤   数量 200 斤    │
│ 地点 城关镇（可改）              │
│ 有效期 7天 ▾   联系 138****5678  │
│ 图片 [+] [+]  加图成交快很多      │
│ [ ] 我承诺信息真实，不收前置费用   │
│          [发布（免费）]           │
└────────────────────────────────┘

屏3 发布结果：
已发布 / 审核中 / 已驳回（原因 + 修改）
浏览 12 · 收藏 3 · 联系 2 · 报名 1
[有人回应时通知] [标记已找到/已招满]
[转发到微信群] [查看我的信息]
""", ["只在标题为空时阻断；未识别字段留空高亮。", "ASR 失败切纯语音发布，页面不弹技术错误。", "发布成功页展示真实浏览、收藏、联系和响应状态，并允许标记已找到/已招满。"]),
        ("6.4 AI 对话页", """
┌────────────────────────────────┐
│ 问问 AI                         │
│ [有没有人收玉米] [镇上有日结活吗] │
│ [帮我发个球局] [怎么防押金骗局]   │
│                                │
│ 你：我是木工，帮我找工作          │
│ AI：找到 3 条，下面卡片来自平台数据 │
│ [招木工2人 260元/天] [查看][报名] │
│ [家具厂长期招木工]   [查看][报名] │
│ 没找到？ [订阅提醒] [帮我发布]     │
│ 也可以帮你挂一个： [帮我发布]      │
│ [按住说话]  输入...        [发送] │
└────────────────────────────────┘
""", ["检索答案必须有卡片，不允许模型自由复述价格。", "最多一个澄清问题，优先用可点选项。", "结果不足三条时追加订阅提醒；零结果直接给发布草稿。"]),
        ("6.5 详情页", """
┌────────────────────────────────┐
│ ‹ 返回        信息详情      分享 │
│ [类目色块/首图]                  │
│ 城关镇土鸡蛋 200斤               │
│ 1.20 元/斤   约 200 斤           │
│ 城关镇 · 同镇 · 2 小时前         │
│ 发布者 老李  发过 4 次 · 0 举报   │
│ 平台提醒：不收定金，不点陌生链接  │
│ 正文...                          │
│ [在线聊] [打电话] [举报]          │
└────────────────────────────────┘
""", ["反诈条固定位置、粗框显示，按类目切换文案。", "联系方式默认脱敏，点击完整号码前先校验配额并记录访问。", "已成交/过期信息可查看但不能继续联系或报名。"]),
        ("6.6 会话页", """
┌────────────────────────────────┐
│ ‹ 返回       和王师傅聊天         │
│ [安全提示：仅限平台内沟通，转账谨慎] │
│                                │
│ 对方：明天来上工，先交 200 押金   │
│ ┌────────────────────────────┐ │
│ │ 反诈提醒：平台不要求前置费用。  │ │
│ │ 如已转账，请立即报警。 [举报]    │ │
│ └────────────────────────────┘ │
│ 输入消息...  [相册] [按住说话] [发送]│
│ [防骗指南]                      │
└────────────────────────────────┘
""", ["首次会话展示一次性提示。", "命中转账/押金/验证码/二维码等词时插入不可关闭反诈卡。", "举报入口在会话内常驻，必要时支持截图。"]),
        ("6.7 我的", """
┌────────────────────────────────┐
│ 我的  老李  城关镇               │
│ [我发布的 3] [待审核 1] [消息 2]  │
│ 在展示  待审  已过期  已成交       │
│ [城关镇土鸡蛋]     [续期]         │
│ [旧打谷机]         [已卖出]       │
│ 我报名的   我收藏的               │
│ 帮助中心   防骗指南   免责声明     │
└────────────────────────────────┘
""", ["状态分组比复杂筛选更适合中老年用户。", "过期信息提供一键续期，续期刷新 bumped_at。", "帮助中心必须包含反诈专页、平台规则和正式免责声明。"]),
        ("6.8 后台审核台", """
┌──────────────────────────────────────────┐
│ 待审 23  举报 4  今日通过率 91%            │
│ [高风险优先] [类目] [乡镇] [批量通过]       │
├──────────────────────────────────────────┤
│ 风险 86  招工  城关镇  2 分钟前             │
│ 原文 / ASR 原声 / 抽取字段 / 命中词 / 历史 │
│ [通过] [修改后通过] [驳回] [下架+封禁+通知] │
└──────────────────────────────────────────┘
""", ["风险分倒序，不按时间堆队列。", "一屏显示判断依据，避免审核员来回切页。", "高危处置组合键必须原子完成并写审计日志。"]),
        ("6.9 系统配置页", """
┌──────────────────────────────────────────┐
│ 系统配置                                   │
│ [账号][发布][审核][反诈][AI][消息][性能]   │
├──────────────────────────────────────────┤
│ 普通用户日发布条数     [ 3 ] 条  范围1-10  │
│ 新用户首发是否人审     [开关]              │
│ 到期提醒提前           [ 4 ] 小时          │
│ AI 提供商              [只读：配置文件]    │
│                              [保存修改]    │
└──────────────────────────────────────────┘
""", ["数字项显示范围并在前端阻止越界。", "危险配置保存前二次确认，变更写 config_change_log。", "editable=false 的基础设施配置只读展示。"]),
    ]
    for title, wire, notes in wireframes:
        doc.add_heading(title, level=2)
        add_code(doc, wire)
        add_bullets(doc, notes)

    doc.add_heading("七、关键状态与异常处理", level=1)
    add_table(doc, ["状态", "用户看到什么", "系统动作", "不可省略的验收点"], [
        ("空列表", "帮你发一条", "提供同类目发布草稿入口", "不能只显示暂无数据"),
        ("无结果搜索", "相近结果 / 订阅提醒 / 帮我发布", "记录 query，生成预填草稿或订阅", "搜索不能让用户空手离开"),
        ("发布后无回应", "已发布，继续完善信息", "提醒补图、补充地点或延长有效期", "不能伪造浏览/响应数据"),
        ("收到回应", "有人收藏/联系/报名了", "写 response 事件并通知发布者", "发布者能标记已找到/已招满/已完成"),
        ("加载中", "骨架屏", "首屏并行请求，缓存优先", "不使用转圈阻塞页面"),
        ("弱网/失败", "网络慢，点这里重试", "保留本地草稿，允许重试", "不丢语音、文字和图片"),
        ("ASR 识别中", "正在识别，可继续操作", "异步轮询或长轮询", "界面不锁死"),
        ("ASR 失败", "先按语音发布，也可打字", "纯语音 Post + Worker 重试", "不弹技术堆栈错误"),
        ("待审核", "审核中，预计 2 小时内", "仅作者可见，进入队列", "他人不可见"),
        ("已驳回/下架", "原因 + 修改/申诉", "保留证据和操作记录", "原因必须具体"),
        ("配额用尽", "还可在某时恢复", "返回剩余时间，预留提升额度", "不让用户误以为系统故障"),
    ], [1500, 2100, 3000, 2760])
    doc.add_heading("八、验收清单与开放问题", level=1)
    add_bullets(doc, [
        "语音发布从进入发布到提交不超过三次主动交互，60 秒语音可在弱网下恢复。",
        "游客浏览正常，但触碰联系方式或发布时才触发登录。",
        "高危内容不上线，中危内容带警示并入人工队列，所有处置可追溯。",
        "零结果搜索与 AI 无结果都能导向发布，不让用户空手离开。",
        "首页搜索、热门词、无结果转发布、发布后响应通知和本镇脉搏均可走通。",
        "所有增长指标都来自真实事件，后台能按乡镇、类目和日期查看漏斗。",
        "招工、二手、约局的类目专属字段和风险提示均由配置驱动。",
        "后台一个审核员可在 30 分钟内处理当日待审与举报队列。",
    ])
    add_callout(doc, "上线前必须确认", "具体县名和覆盖乡镇、作物/工种/地名词库、先发后审合规路径、运营审核人力、服务器与 AI 月预算、语音与聊天留存期限、置顶收费启用时间。", fill="FFF8E8", color=GOLD)
    return doc


def add_design_doc():
    doc = configure_document(Document(), "单人 AI 落地方案")
    add_title_block(doc, "县域生活信息交换服务平台", "适合一人 AI 开发落地的设计方案", "F-ARCH-20260902")
    add_callout(doc, "推荐结论", "采用模块化单体 + 纵向切片开发：Taro/React 小程序与 H5 同构，NestJS/TypeScript 后端，MySQL + Redis + OSS/CDN，所有 AI 能力经过 AiGateway。首期不引入微服务、ES、独立向量库和复杂 RBAC。")
    doc.add_heading("阅读导航", level=1)
    add_bullets(doc, [
        "第一至三章：范围、技术选型和架构边界。",
        "第四至六章：数据模型、接口和关键业务工作流。",
        "第七至九章：AI 网关、配置中心、安全治理和运维。",
        "第十至十二章：10 周排期、测试、成本预算和扩展窗口。",
    ])

    doc.add_heading("一、单人开发的产品化原则", level=1)
    add_bullets(doc, [
        "先闭环再丰富：先确保一个用户能发布、一名雇主能联系、运营能处置。",
        "A/B/C 先做真实信息交换闭环：让用户更容易查、发了有回音、每天有理由回来。",
        "共用骨架：列表、详情、发布、审核、搜索都围绕统一 Post，不为四个模块复制页面。",
        "规则替代模型：价格、数量、日期、电话、地名和品种优先正则/词库，模型只处理标题和正文清理。",
        "同步返回、异步补全：发布接口先落草稿和语音地址，ASR、抽取、审核、向量化进队列。",
        "所有外部能力可替换：AI、ASR、TTS、图片审核、支付都挂在接口后，业务代码不直接 import SDK。",
        "可观测优先：每个关键动作都有 request_id、post_id、user_id、耗时、状态和成本记录。",
    ])
    add_table(doc, ["范围", "首期交付", "明确延后"], [
        ("P0", "登录、乡镇、语音发布、统一 Post、审核、浏览、搜索、AI 双出口、私聊、反诈、我的、后台治理", "无"),
        ("P1", "草稿恢复、相似去重、分享海报、TTS、行为展示、订阅提醒、每日摘要、配置高级能力", "可在 MVP 后半段插入"),
        ("P2", "互评、钱包、推广收费、担保交易、会员、多县多租户", "不影响首期上线"),
    ], [1200, 4300, 3860])

    doc.add_heading("二、技术栈与部署形态", level=1)
    add_table(doc, ["层", "选择", "为什么适合一人开发", "迁移窗口"], [
        ("客户端", "Taro + React + TypeScript", "双端同构、生态成熟、AI 代码生成质量高", "platform-adapter 隔离微信能力"),
        ("后端", "NestJS + TypeScript", "模块边界清晰、校验/鉴权/队列容易组织", "DAO 分 readDb/writeDb"),
        ("数据库", "MySQL 8", "统一 Post、全文 ngram、百万行级够用", "预留读写分离"),
        ("缓存/队列", "Redis + Redis Stream", "一个组件同时做缓存、限流、会话和异步队列", "后续可替换消息系统"),
        ("文件", "OSS + CDN", "图片和语音不压垮应用服务器", "签名 URL 可换供应商"),
        ("后台", "Vue 3 + Element Plus 或 Nest SSR", "访问量低，优先可维护", "不引入独立权限服务"),
        ("部署", "Docker Compose + Nginx", "2C4G 单机可平移，命令少", "第三阶段再考虑 K8s"),
    ], [1200, 1900, 3100, 3160])
    add_code(doc, """
docker compose services
├─ app-api       NestJS API + scheduler
├─ app-worker    Redis Stream consumers (asr/extract/audit/notify/embed)
├─ admin         管理端（可与 app-api 同镜像，独立端口）
├─ mysql         MySQL 8，定时备份 + binlog
├─ redis         Redis 7，maxmemory 512MB + LRU
└─ nginx         HTTPS、gzip/brotli、管理端 IP 白名单
""")

    doc.add_heading("三、代码仓库结构", level=1)
    add_code(doc, """
repo/
├─ apps/
│  ├─ client/                  # Taro 小程序 + H5
│  ├─ admin/                   # 运营后台
│  └─ api/                     # NestJS API
├─ packages/
│  ├─ domain/                  # Post/User/Deal/Config 类型与状态机
│  ├─ shared/                  # DTO、错误码、日志、时间与枚举
│  ├─ platform-adapter/        # record/login/location/subscribe
│  └─ ai-contracts/            # AiGateway 接口与供应商适配器
├─ workers/
│  ├─ asr.worker.ts
│  ├─ post-extract.worker.ts
│  ├─ audit.worker.ts
│  ├─ notify.worker.ts
│  └─ embedding.worker.ts
├─ database/
│  ├─ migrations/
│  ├─ seeds/                   # 类目、行政区划、词库、FAQ、风险词
│  └─ sql/
├─ infra/
│  ├─ docker-compose.yml
│  ├─ nginx.conf
│  └─ backup.ps1
└─ docs/
   ├─ api.md
   ├─ runbook.md
   └─ eval/
""")
    add_callout(doc, "一人开发边界", "不要拆成多个仓库，也不要把每个业务模块做成独立服务。单人最怕的是跨仓库联调和重复 DTO；模块化单体足以承载单县 MVP。")

    add_growth_architecture(doc)

    doc.add_heading("四、数据模型与状态", level=1)
    add_table(doc, ["表", "关键字段", "首期职责", "索引/约束"], [
        ("users", "id, openid, phone_enc, town_code, verified_at, role", "账号、实名溯源、角色", "openid/phone 唯一，敏感字段加密"),
        ("posts", "统一 Post 字段 + ext JSON", "四模块内容、状态、时效", "status+category+town_code+bumped_at"),
        ("post_media", "post_id, type, url, sha256, lqip", "图片/语音元数据", "哈希查重，签名 URL"),
        ("deals", "post_id, buyer_id, seller_id, state", "记录联系、成交意向，首期不收款", "CONTACTED -> AGREED -> CLOSED"),
        ("conversations/messages", "conversation_id, sender_id, content, risk_flags", "私聊证据链和反诈", "按会话+时间分页"),
        ("audit_cases/logs", "post_id, risk_score, action, operator_id", "机器/人工审核留痕", "不可删除，按时间检索"),
        ("config/config_change_log", "key, value, type, min/max, updated_by", "67 项可配置参数和审计", "key 唯一，变更必写日志"),
        ("dictionaries", "kind, term, normalized, version", "作物、工种、地名、同义词、风险词", "kind+term 唯一"),
        ("events", "event_name, actor, payload, created_at", "指标、数据飞轮、回放", "按事件和日期分区式归档"),
    ], [1500, 3000, 3000, 1860])
    add_code(doc, """
type PostStatus = 'draft' | 'pending_review' | 'published'
  | 'closed' | 'expired' | 'removed' | 'rejected';

type PostDraft = {
  category?: Category;
  subCategory?: string;
  title?: string;
  bodyText?: string;
  price?: number;
  priceUnit?: string;
  quantity?: number;
  quantityUnit?: string;
  locationCode?: string;
  validUntil?: string;
  ext: Record<string, unknown>;
  missingFields: string[];
  warnings: string[];
};

type GrowthEvent = {
  eventName: 'search_submitted' | 'search_zero_result' | 'search_result_clicked'
    | 'draft_created' | 'post_published' | 'response_received'
    | 'post_completed' | 'digest_opened' | 'expiry_reminded' | 'post_renewed';
  actorId?: number;
  postId?: number;
  townCode: string;
  payload: Record<string, unknown>;
  createdAt: string;
};
""")

    doc.add_heading("五、核心接口设计", level=1)
    add_table(doc, ["接口", "方法", "作用", "关键返回"], [
        ("/auth/wechat-login", "POST", "小程序登录、绑定手机号", "access_token, user_profile"),
        ("/home/bootstrap", "GET", "首页一次请求返回配置、类目、首屏列表", "config_version, categories, posts"),
        ("/media/presign", "POST", "获取 OSS 上传签名", "upload_url, object_key, expires_at"),
        ("/posts/drafts", "POST/PATCH", "创建或更新草稿", "draft_id, revision"),
        ("/posts/{id}/publish", "POST", "提交审核和上线流程", "post_id, status, reason"),
        ("/posts", "GET", "类目/乡镇/排序/游标列表", "items, next_cursor"),
        ("/search", "GET/POST", "关键词或语音搜索", "items, fallback_action"),
        ("/search/suggestions", "GET", "返回本镇热搜、急需和最近搜索词", "items, source, expires_at"),
        ("/subscriptions", "POST/DELETE", "订阅关键词、乡镇和类目，支持退订", "subscription_id, status"),
        ("/posts/{id}/responses", "GET", "返回浏览、收藏、联系、报名和响应状态", "counts, latest_events"),
        ("/posts/{id}/status", "POST", "标记已找到、已招满或已完成", "post_id, status, event_id"),
        ("/events", "POST", "记录搜索、发布、响应和摘要行为事件", "event_id"),
        ("/ai/chat/stream", "POST SSE", "AI 对话和双出口", "events: text/cards/actions"),
        ("/conversations/{id}/messages", "GET/POST", "长轮询消息和发送", "messages, next_cursor"),
        ("/reports", "POST", "举报、截图和证据", "case_id"),
        ("/admin/audit-cases", "GET/POST", "审核队列和处置", "case, action_log"),
        ("/admin/config", "GET/PATCH", "配置读取和修改", "version, change_log_id"),
    ], [2150, 1250, 3250, 2710])
    add_callout(doc, "接口约束", "所有写接口带 Idempotency-Key；所有异步任务带 trace_id 和 attempt；列表接口只返回卡片字段；联系方式必须通过单独的授权接口获取。", fill="FFF8E8", color=GOLD)

    doc.add_heading("六、关键业务工作流", level=1)
    doc.add_heading("6.1 语音发布", level=2)
    add_code(doc, """
客户端录音 -> OSS 直传 -> POST /posts/drafts
                     │
                     └-> Redis Stream: asr.request
                                   -> ASR / 8s timeout
                                   -> 规则归一化（价格/数量/时间/地名）
                                   -> AiGateway.extractPostFields
                                   -> PostDraft + missingFields + warnings
                                   -> 客户端轮询草稿
用户确认 -> POST /posts/{id}/publish
         -> L1 风险词 -> L2 图片机审 -> L3 行为风控
         -> published / pending_review / rejected
""")
    add_bullets(doc, [
        "同步阶段只做签名、草稿落库和状态返回，目标 < 500ms。",
        "ASR > 8 秒或服务故障不阻塞发布，保留 raw_voice_url，后台补齐 raw_asr_text。",
        "抽取结果必须返回 missingFields 和 warnings，不把不确定性伪装成已确认值。",
        "用户确认前不得静默发布，用户修改 diff 写入 events 作为模型评测数据。",
    ])
    doc.add_heading("6.2 AI 对话双出口", level=2)
    add_code(doc, """
query -> L0 规则/同义词 -> L1 嵌入意图 -> L2 大模型兜底
      -> 槽位抽取(category/town/price/time)
      -> SQL 硬过滤(status=published, valid_until>now)
      -> MySQL ngram + 内存向量双路召回
      -> RRF 融合 + 距离/时新度/quality_score 重排
      -> 结果卡片 + 联系/报名 + 帮我发布 + 订阅提醒
""")
    add_bullets(doc, [
        "系统最多问一个澄清问题，而且必须用按钮选项，不开放多轮盘问。",
        "没有结果时不返回‘没找到’，直接生成待确认发布草稿。",
        "平台规则和反诈答案来自 FAQ/模板，不让模型自由生成。",
        "模型不可用时降级为关键词搜索 + 发布入口，不显示模型错误。",
    ])
    doc.add_heading("6.3 审核与反诈", level=2)
    add_code(doc, """
发布/私聊/举报事件
  -> 规则词库与正则
  -> 图片安全 API（只在有图时调用）
  -> 行为风控：新号、频次、重复、举报次数、金额阈值
  -> 自动通过 / 待审 / 驳回 / 临时封禁
  -> 通知发布者、联系过该信息的用户、运营值班人
""")

    doc.add_heading("七、AI 能力网关", level=1)
    add_code(doc, """
export interface AiGateway {
  classifyIntent(text: string): Promise<IntentResult>;
  extractPostFields(text: string, category?: Category): Promise<PostDraft>;
  embed(text: string, modelVersion?: string): Promise<number[]>;
  chat(input: ChatInput): AsyncIterable<ChatEvent>;
  transcribe(audioUrl: string): Promise<{ text: string; confidence: number }>;
  synthesize(text: string): Promise<{ audioUrl: string; cacheKey: string }>;
  moderate(input: { text?: string; imageUrl?: string }): Promise<RiskResult>;
}
""")
    add_table(doc, ["能力", "首期实现", "降级", "观测指标"], [
        ("Intent", "规则 -> 小嵌入中心向量", "关键词搜索", "路由分布、准确率"),
        ("Extract", "规则归一 + 小模型 JSON Schema", "纯语音/手填", "字段完整率、修改率"),
        ("ASR", "云 ASR，8 秒超时", "raw voice 发布 + 后台补齐", "P95、失败率、分钟成本"),
        ("Chat", "SSE，最近 3 轮上下文", "搜索 + 快捷入口", "首字延迟、调用占比"),
        ("TTS", "按 post_id 缓存到 OSS", "隐藏听一听按钮", "缓存命中率、成本"),
        ("Moderation", "规则 + 图片云审 + 行为", "人工队列", "误报率、审核 SLA"),
    ], [1500, 3100, 2500, 2260])
    add_callout(doc, "模型切换规则", "业务代码只依赖 AiGateway。供应商、模型名、超时和 fallback 写在环境变量/配置文件；向量表保存 dim 和 model_version，支持新旧向量并存、双写、回填、切换和回滚。")

    doc.add_heading("八、配置中心与可扩展性", level=1)
    add_table(doc, ["类别", "数量", "实现方式", "首期验收"], [
        ("账号与权限", "7", "config 表 + ConfigService", "限额、查看次数、奖励实时生效"),
        ("发布与有效期", "8", "按类目读取默认值", "valid_until、图片上限、提醒时间可改"),
        ("审核规则", "12", "audit_rules + 60 秒刷新", "高/中危词、金额阈值可改"),
        ("反诈拦截", "5", "关键词、URL、微信号开关", "私聊命中规则立即生效"),
        ("AI/消息/性能", "35", "环境变量 + config + Redis", "模板、缓存和成本参数有来源"),
    ], [1800, 1000, 3300, 3260])
    add_code(doc, """
interface ConfigService {
  getInt(key: string): number;
  getFloat(key: string): number;
  getBool(key: string): boolean;
  getString(key: string): string;
  getJson<T>(key: string): T;
  set(key: string, value: string, adminId: number): Promise<void>;
  reload(): Promise<void>;
}
""")
    add_bullets(doc, [
        "业务代码禁止出现日发布条数、有效期、审核阈值等魔法数字。",
        "配置保存前按 min_value/max_value 校验；危险项二次确认；每次修改写旧值、新值、操作者和 IP。",
        "类目与 ext 字段模板配置化，新增类目时不改后端表结构。",
        "Feature Flag 首期存在但默认关闭：payment.enabled、escrow.enabled、promotion.enabled、member.enabled。",
    ])

    doc.add_heading("九、安全、隐私与合规门禁", level=1)
    add_table(doc, ["区域", "必须实现", "上线门禁"], [
        ("身份与联系方式", "手机号脱敏、完整查看授权、访问日志、敏感字段加密", "抽样检查访问日志和加密列"),
        ("上传", "MIME/大小校验、OSS 签名 URL、图片审核、哈希", "上传恶意类型被拒绝"),
        ("管理端", "独立域名/端口、IP 白名单或 VPN、2FA、角色枚举中间件", "公网扫描面与审计日志确认"),
        ("反诈", "发布/详情/私聊/事后四层拦截、举报、主动预警", "模拟押金、二维码、赌博话术"),
        ("隐私", "最小采集、位置精度到乡镇、导出与注销、留存清理", "律师确认语音/聊天留存期"),
        ("合规", "免责声明、招工风险、打牌承诺、平台非交易方文案", "网信/市场监管/律师逐项确认"),
    ], [1700, 4700, 2960])
    add_callout(doc, "不可绕过的业务门", "如果当地主管部门要求先审后发，就启用人工值班并把非值班时段标为待审；不能只靠技术默认值替代合规结论。", fill="FFF8E8", color=GOLD)

    doc.add_heading("十、10 周单人开发排期", level=1)
    add_table(doc, ["周次", "交付目标", "具体任务", "完成标准"], [
        ("1", "骨架与适配层", "仓库、Docker、登录、platform-adapter、统一错误码、日志、Post migration", "小程序/H5 可登录，CI 可跑"),
        ("2", "账号与位置", "手机号/身份补全、乡镇选择、权限、ConfigService、QuotaService", "游客/用户/运营权限可验证"),
        ("3", "语音与草稿", "长按录音、OSS 签名、ASR worker、草稿恢复、纯语音降级", "60 秒语音能落 OSS，失败不丢稿"),
        ("4", "发布与审核", "抽取、三层审核、状态机、通知、我的发布", "低/中/高风险三条路径可回归"),
        ("5", "首页与列表", "bootstrap、逛一逛、游标分页、缓存、详情", "首页首屏和详情闭环"),
        ("6", "四模块差异", "ext 表单、农产品、招工任务、二手、约局", "四类信息可发布/浏览/联系"),
        ("7", "搜索与 AI", "三级意图、混合检索、SSE、双出口、零结果兜底、热搜词", "三组场景 query 通过验收，零结果可转发布"),
        ("8", "私聊与反诈", "长轮询、未读、四层反诈、举报、封禁止损", "押金/二维码/赌博脚本可拦截"),
        ("9", "后台与公众号", "审核台、配置台、增长漏斗、摘要、菜单、周报模板", "运营一人可完成日常治理和召回"),
        ("10", "压测与试点", "弱网、兼容、安全、备份恢复、单乡镇灰度", "上线门禁通过，回滚可执行"),
    ], [850, 1900, 4250, 2360])
    add_callout(doc, "排期纪律", "每周只认一个可演示闭环。AI 辅助生成代码必须在同一周完成类型检查、接口测试和浏览器/真机冒烟，不积压到最后一周。")

    doc.add_heading("十一、测试与验收策略", level=1)
    add_table(doc, ["测试层", "覆盖内容", "最低要求"], [
        ("单元测试", "状态机、QuotaService、ConfigService、规则匹配、中文数字归一化", "核心规则覆盖率 >= 90%"),
        ("接口测试", "登录、发布、审核、搜索、AI 降级、消息、举报", "每条 P0 至少一条成功 + 一条失败"),
        ("流程测试", "语音发布、报名、满员、续期、封禁止损", "Playwright/真机脚本可重复运行"),
        ("增长闭环测试", "搜索零结果、query 转草稿、响应通知、状态更新、摘要打开", "A/B/C 四段链路事件完整且不重复"),
        ("AI 评测", "100-200 条 query、Recall@5、MRR、字段完整率、修改率", "每次换 Prompt/模型自动跑"),
        ("弱网测试", "上传中断、ASR 超时、SSE 断线、离线草稿", "4G 慢网下可恢复"),
        ("安全测试", "鉴权、越权、导出、上传、限流、SQL 注入", "管理端与联系方式重点验收"),
    ], [1500, 5000, 2860])
    add_code(doc, """
验收脚本最小集：
1. 游客浏览 -> 点击联系方式 -> 登录 -> 身份补全 -> 查看并写日志
2. 长按录音 -> ASR 正常 -> 修改字段 -> 发布 -> 自动通过
3. ASR 超时 -> 纯语音发布 -> Worker 补齐文本
4. 招工缺日薪 -> 高亮提示；押金词 -> 警示/待审；赌博词 -> 拦截封禁
5. AI “找木工” -> 三类结果卡片 -> 报名/联系 -> 双出口发布草稿
6. 举报三次 -> 临时封禁 -> 批量通知近 7 天会话对象
7. 搜索“收玉米”无结果 -> 相近词/订阅提醒/帮我发布 -> 草稿创建事件
8. 发布后被收藏/联系 -> 发布者收到通知 -> 标记已找到 -> 摘要聚合
""")

    doc.add_heading("十二、性能、成本与运维", level=1)
    add_table(doc, ["预算项", "目标", "实现手段", "监控"], [
        ("首页首屏", "4G < 1.5s", "bootstrap 合并接口、60 秒缓存、骨架屏", "P50/P95、缓存命中"),
        ("列表接口", "P95 < 300ms", "结构化过滤、游标分页、列表字段裁剪", "慢查询 > 200ms"),
        ("AI 首字", "< 1.5s", "规则/L1 路由、SSE、上下文只带 3 轮", "首字延迟、路由占比"),
        ("发布落库", "< 500ms", "OSS 直传、异步 ASR/审核", "写入耗时、队列长度"),
        ("单条信息成本", "<= 0.02 元", "规则替代、TTS/摘要缓存、模型分层", "按日 token/ASR/图片审核成本"),
        ("备份恢复", "每日全量 + binlog", "定时任务、每月恢复演练", "备份成功率、恢复时长"),
    ], [1600, 1600, 4000, 2160])
    add_bullets(doc, [
        "Redis 仅做缓存/队列/限流，不把唯一业务数据放在 Redis。",
        "所有文件走 OSS/CDN，应用服务器不存用户图片和语音。",
        "消息、审核、向量化、统计全部异步；热点信息和空结果都做短 TTL 缓存。",
        "先单乡镇试点，运营手动撮合，收集用户修改 diff 和搜索无结果 query，再调 Prompt/词库。",
    ])
    doc.add_heading("十三、AI 辅助开发工作流", level=1)
    add_table(doc, ["阶段", "AI 用法", "人工必须把关"], [
        ("需求拆解", "把 FR 表转换为 DTO、接口、验收用例", "确认业务边界、风险词和合规文案"),
        ("编码", "生成样板、迁移、CRUD、测试数据", "审查状态机、鉴权、幂等、错误处理"),
        ("调试", "基于日志和失败用例定位问题", "不接受 AI 猜测的线上修复"),
        ("评测", "批量运行 query、生成差异报告", "人工确认真实结果和反诈答案"),
        ("发布", "生成变更说明、回滚命令、运维清单", "执行灰度、备份、验收和回滚"),
    ], [1600, 4200, 3560])
    add_callout(doc, "代码审查底线", "AI 生成的每个写接口都必须检查：鉴权、越权、幂等、限流、输入校验、日志、事务边界、敏感数据输出和失败重试。")
    doc.add_heading("十四、阶段二与阶段三迁移窗口", level=1)
    add_table(doc, ["能力", "首期预留", "启用时动作", "不应改动"], [
        ("置顶收费", "is_top/top_expire_at/weight_boost + promotion flag", "新增 PaymentProvider 实现", "Post 排序接口"),
        ("担保交易", "deal + EscrowProvider + 订单页骨架", "接入持牌机构，启用 escrow flag", "撮合与消息链路"),
        ("更强模型", "AiGateway、model_version、灰度哈希", "改配置、双写向量、5% 灰度", "业务调用方"),
        ("向量库", "MySQL BLOB + 内存索引", "数据 > 20 万条且服务器升级后迁 Qdrant/同类", "检索接口契约"),
        ("多县多租户", "location_code、tenant_id 预留", "新增租户边界和配置隔离", "统一 Post 领域模型"),
    ], [1700, 3500, 2600, 1560])
    add_callout(doc, "最终交付判断", "当一个人能在本地启动、跑通一条语音发布、完成审核、模拟私聊反诈、查看指标并从备份恢复时，MVP 才算完成。只有页面能打开，不算完成。", fill="FFF8E8", color=GOLD)
    return doc


if __name__ == "__main__":
    function_doc = add_function_doc()
    design_doc = add_design_doc()
    function_path = OUT / "县域生活信息交换服务平台-功能点明细与低保真解析-V1.1-ABC融合.docx"
    design_path = OUT / "县域生活信息交换服务平台-单人AI开发落地设计方案-V1.1-ABC融合.docx"
    function_doc.save(function_path)
    design_doc.save(design_path)
    print(function_path)
    print(design_path)
